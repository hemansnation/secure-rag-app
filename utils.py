import warnings
warnings.filterwarnings("ignore", category=FutureWarning)

import logging
from cryptography.fernet import Fernet
import os
from PyPDF2 import PdfReader
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from dotenv import load_dotenv
from langchain_core.documents import Document


load_dotenv()

# Setup logging for audits
logging.basicConfig(filename='rag_audit.log', level=logging.INFO, format='%(asctime)s - %(message)s')

# Generate or load encryption key (for data-at-rest)
key_file = 'secret.key'
if not os.path.exists(key_file):
    key = Fernet.generate_key()
    with open(key_file, 'wb') as f:
        f.write(key)
with open(key_file, 'rb') as f:
    key = f.read()
cipher = Fernet(key)

def load_document(file_path):
    """Load and extract text from PDF or Doc securely."""
    if file_path.endswith('.pdf'):
        reader = PdfReader(file_path)
        text = ''.join(page.extract_text() for page in reader.pages)
    elif file_path.endswith('.docx'):
        doc = Document(file_path)
        text = '\n'.join(para.text for para in doc.paragraphs)
    else:
        raise ValueError("Unsupported file type")
    
    # Encrypt text for storage (decrypt on use)
    encrypted_text = cipher.encrypt(text.encode())
    logging.info(f"Loaded and encrypted document: {file_path}")
    return encrypted_text.decode()  # Return encrypted as string for now

def decrypt_text(encrypted_text):
    return cipher.decrypt(encrypted_text.encode()).decode()



def chunk_text(text, chunk_size=512, overlap=50):
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=overlap)
    chunks = splitter.split_text(text)
    # Wrap as LangChain Documents with metadata for citations
    docs = []
    for i, chunk in enumerate(chunks):
        docs.append(Document(
            page_content=chunk,
            metadata={"chunk_id": i+1, "source": "uploaded_document"}  # Can add page_num if from PDF
        ))
    return docs